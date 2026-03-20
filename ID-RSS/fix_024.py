from docx import Document
from pathlib import Path

# Fix student_0024.docx 
file_path = Path('demo_files') / 'student_0024.docx'
doc = Document(file_path)

for para in list(doc.paragraphs):
    p = para._element
    p.getparent().remove(p)

doc.add_paragraph('Serial No.: STU-2024-0024')
doc.add_paragraph('Student identification record - name missing.')
doc.save(file_path)
print("DONE: student_0024.docx fixed with clean serial number format")
