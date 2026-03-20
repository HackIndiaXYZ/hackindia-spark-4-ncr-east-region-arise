#!/usr/bin/env python3
"""Fix demo files to show proper status badges without file corruption"""

from docx import Document
from pathlib import Path
import subprocess

# Restore ALL modified files from git to get clean versions
files_to_restore = [
    'demo_files/student_0001.docx',
    'demo_files/student_0004.docx', 
    'demo_files/student_0009.docx',
    'demo_files/student_0021.docx',
    'demo_files/student_0024.docx'
]

print("=" * 60)
print("RESTORING CLEAN VERSIONS FROM GIT...")
print("=" * 60)

for f in files_to_restore:
    result = subprocess.run(['git', 'checkout', f], capture_output=True, text=True)
    if result.returncode == 0:
        print(f"✓ Restored {f}")
    else:
        print(f"✗ Failed: {result.stderr}")

print("\n" + "=" * 60)
print("MODIFYING WITH VALID .docx CONTENT...")
print("=" * 60 + "\n")

demo_dir = Path('demo_files')

modifications = {
    'student_0001.docx': {'type': 'empty', 'desc': 'NO extract data'},
    'student_0004.docx': {'type': 'empty', 'desc': 'NO extract data'}, 
    'student_0009.docx': {'type': 'empty', 'desc': 'NO extract data'},
    'student_0021.docx': {'type': 'missing_serial', 'desc': 'Only Name present'},
    'student_0024.docx': {'type': 'missing_name', 'desc': 'Only Serial No. present'},
}

for filename, config in modifications.items():
    file_path = demo_dir / filename
    
    try:
        doc = Document(file_path)
        
        # Clear all existing paragraphs
        for para in list(doc.paragraphs):
            p = para._element
            p.getparent().remove(p)
        
        if config['type'] == 'empty':
            doc.add_paragraph('Academic records - No identification data present')
            print(f"✓ {filename}: {config['desc']} (Orange badge)")
            
        elif config['type'] == 'missing_serial':
            doc.add_paragraph('Name: Student from India')
            doc.add_paragraph('Record incomplete - no serial number.')
            print(f"✓ {filename}: {config['desc']} (Serial No. missing badge)")
            
        elif config['type'] == 'missing_name':
            doc.add_paragraph('Serial No.: STU-2024-0024')
            doc.add_paragraph('Student identification record - name missing.')
            print(f"✓ {filename}: {config['desc']} (Name missing badge)")
        
        doc.save(file_path)
        print(f"   → Saved as valid .docx file\n")
        
    except Exception as e:
        print(f"✗ ERROR on {filename}: {e}\n")

print("=" * 60)
print("✓ ALL FILES FIXED!")
print("Refresh browser at http://127.0.0.1:5000")
print("=" * 60)
